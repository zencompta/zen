"""
Service de génération de rapports avec templates personnalisables
Gère la création de rapports professionnels avec branding client et formats multiples
"""

import os
import json
import uuid
from datetime import datetime
from typing import Dict, List, Optional, Any
from dataclasses import dataclass
from pathlib import Path
import base64
from io import BytesIO

# Imports pour la génération de documents
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import Color, HexColor
from reportlab.lib import colors
from reportlab.graphics.shapes import Drawing
from reportlab.graphics.charts.linecharts import HorizontalLineChart
from reportlab.graphics.charts.barcharts import VerticalBarChart

# Imports pour Excel
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.chart import LineChart, BarChart, Reference
from openpyxl.drawing.image import Image as ExcelImage

# Imports pour Word
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

@dataclass
class BrandingConfig:
    """Configuration du branding client"""
    logo_path: Optional[str] = None
    primary_color: str = "#1f2937"  # Couleur principale
    secondary_color: str = "#3b82f6"  # Couleur secondaire
    accent_color: str = "#10b981"  # Couleur d'accent
    company_name: str = ""
    company_address: str = ""
    company_phone: str = ""
    company_email: str = ""
    header_text: str = ""
    footer_text: str = ""
    font_family: str = "Arial"

@dataclass
class ReportTemplate:
    """Template de rapport personnalisable"""
    template_id: str
    name: str
    description: str
    sector: str  # Secteur d'activité
    template_type: str  # financial, audit, compliance, etc.
    sections: List[Dict[str, Any]]
    branding: BrandingConfig
    format_options: List[str]  # PDF, Word, Excel, PowerPoint

class ReportGeneratorService:
    """Service de génération de rapports professionnels"""
    
    def __init__(self):
        self.templates_dir = Path("templates")
        self.output_dir = Path("generated_reports")
        self.templates_dir.mkdir(exist_ok=True)
        self.output_dir.mkdir(exist_ok=True)
        
        # Templates par secteur d'activité
        self.sector_templates = {
            "banking": "Modèles bancaires et financiers",
            "insurance": "Modèles d'assurance",
            "retail": "Modèles de commerce de détail",
            "manufacturing": "Modèles industriels",
            "healthcare": "Modèles de santé",
            "technology": "Modèles technologiques",
            "real_estate": "Modèles immobiliers",
            "education": "Modèles éducatifs",
            "government": "Modèles gouvernementaux",
            "nonprofit": "Modèles associatifs"
        }
        
        # Initialiser les templates par défaut
        self._initialize_default_templates()
    
    def _initialize_default_templates(self):
        """Initialise les templates par défaut pour chaque secteur"""
        default_templates = [
            {
                "template_id": "audit_financial_banking",
                "name": "Rapport d'Audit Financier - Banque",
                "description": "Template d'audit financier spécialisé pour le secteur bancaire",
                "sector": "banking",
                "template_type": "audit",
                "sections": [
                    {"type": "cover", "title": "Page de couverture", "required": True},
                    {"type": "executive_summary", "title": "Résumé exécutif", "required": True},
                    {"type": "financial_position", "title": "Position financière", "required": True},
                    {"type": "risk_assessment", "title": "Évaluation des risques", "required": True},
                    {"type": "regulatory_compliance", "title": "Conformité réglementaire", "required": True},
                    {"type": "recommendations", "title": "Recommandations", "required": True},
                    {"type": "appendices", "title": "Annexes", "required": False}
                ],
                "format_options": ["PDF", "Word", "Excel"]
            },
            {
                "template_id": "compliance_ifrs_general",
                "name": "Rapport de Conformité IFRS",
                "description": "Template de conformité aux normes IFRS",
                "sector": "general",
                "template_type": "compliance",
                "sections": [
                    {"type": "cover", "title": "Page de couverture", "required": True},
                    {"type": "scope", "title": "Périmètre d'audit", "required": True},
                    {"type": "ifrs_compliance", "title": "Conformité IFRS", "required": True},
                    {"type": "material_issues", "title": "Points significatifs", "required": True},
                    {"type": "conclusion", "title": "Conclusion", "required": True}
                ],
                "format_options": ["PDF", "Word"]
            }
        ]
        
        for template_data in default_templates:
            template_file = self.templates_dir / f"{template_data['template_id']}.json"
            if not template_file.exists():
                with open(template_file, 'w', encoding='utf-8') as f:
                    json.dump(template_data, f, indent=2, ensure_ascii=False)
    
    def get_templates_by_sector(self, sector: str) -> List[Dict[str, Any]]:
        """Récupère les templates disponibles pour un secteur"""
        templates = []
        for template_file in self.templates_dir.glob("*.json"):
            try:
                with open(template_file, 'r', encoding='utf-8') as f:
                    template_data = json.load(f)
                    if template_data.get('sector') == sector or sector == 'all':
                        templates.append(template_data)
            except Exception as e:
                print(f"Erreur lors du chargement du template {template_file}: {e}")
        return templates
    
    def create_custom_template(self, template_data: Dict[str, Any]) -> str:
        """Crée un template personnalisé"""
        template_id = f"custom_{uuid.uuid4().hex[:8]}"
        template_data['template_id'] = template_id
        template_data['created_at'] = datetime.now().isoformat()
        
        template_file = self.templates_dir / f"{template_id}.json"
        with open(template_file, 'w', encoding='utf-8') as f:
            json.dump(template_data, f, indent=2, ensure_ascii=False)
        
        return template_id
    
    def generate_report(self, 
                       template_id: str, 
                       data: Dict[str, Any], 
                       branding: BrandingConfig,
                       output_format: str = "PDF") -> str:
        """Génère un rapport basé sur un template"""
        
        # Charger le template
        template_file = self.templates_dir / f"{template_id}.json"
        if not template_file.exists():
            raise ValueError(f"Template {template_id} non trouvé")
        
        with open(template_file, 'r', encoding='utf-8') as f:
            template_data = json.load(f)
        
        # Générer le rapport selon le format demandé
        if output_format.upper() == "PDF":
            return self._generate_pdf_report(template_data, data, branding)
        elif output_format.upper() == "WORD":
            return self._generate_word_report(template_data, data, branding)
        elif output_format.upper() == "EXCEL":
            return self._generate_excel_report(template_data, data, branding)
        else:
            raise ValueError(f"Format {output_format} non supporté")
    
    def _generate_pdf_report(self, template: Dict[str, Any], data: Dict[str, Any], branding: BrandingConfig) -> str:
        """Génère un rapport PDF"""
        report_id = f"report_{uuid.uuid4().hex[:8]}"
        output_file = self.output_dir / f"{report_id}.pdf"
        
        # Créer le document PDF
        doc = SimpleDocTemplate(str(output_file), pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Style personnalisé avec branding
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontName=branding.font_family,
            fontSize=24,
            textColor=HexColor(branding.primary_color),
            spaceAfter=30
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontName=branding.font_family,
            fontSize=16,
            textColor=HexColor(branding.secondary_color),
            spaceAfter=12
        )
        
        # En-tête avec logo et informations entreprise
        if branding.logo_path and os.path.exists(branding.logo_path):
            logo = Image(branding.logo_path, width=2*inch, height=1*inch)
            story.append(logo)
            story.append(Spacer(1, 12))
        
        # Titre du rapport
        story.append(Paragraph(template['name'], title_style))
        story.append(Spacer(1, 20))
        
        # Informations de l'entreprise
        if branding.company_name:
            company_info = f"""
            <b>{branding.company_name}</b><br/>
            {branding.company_address}<br/>
            Tél: {branding.company_phone}<br/>
            Email: {branding.company_email}
            """
            story.append(Paragraph(company_info, styles['Normal']))
            story.append(Spacer(1, 20))
        
        # Générer les sections du rapport
        for section in template['sections']:
            if section['required'] or section['type'] in data:
                story.append(Paragraph(section['title'], heading_style))
                
                # Contenu de la section
                section_content = data.get(section['type'], "Contenu à compléter")
                if isinstance(section_content, dict):
                    # Traiter les données structurées
                    for key, value in section_content.items():
                        story.append(Paragraph(f"<b>{key}:</b> {value}", styles['Normal']))
                elif isinstance(section_content, list):
                    # Traiter les listes
                    for item in section_content:
                        story.append(Paragraph(f"• {item}", styles['Normal']))
                else:
                    story.append(Paragraph(str(section_content), styles['Normal']))
                
                story.append(Spacer(1, 15))
        
        # Pied de page
        if branding.footer_text:
            story.append(Spacer(1, 30))
            story.append(Paragraph(branding.footer_text, styles['Normal']))
        
        # Construire le PDF
        doc.build(story)
        
        return str(output_file)
    
    def _generate_word_report(self, template: Dict[str, Any], data: Dict[str, Any], branding: BrandingConfig) -> str:
        """Génère un rapport Word"""
        report_id = f"report_{uuid.uuid4().hex[:8]}"
        output_file = self.output_dir / f"{report_id}.docx"
        
        # Créer le document Word
        doc = Document()
        
        # Ajouter le logo si disponible
        if branding.logo_path and os.path.exists(branding.logo_path):
            doc.add_picture(branding.logo_path, width=Inches(2))
        
        # Titre du rapport
        title = doc.add_heading(template['name'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Informations de l'entreprise
        if branding.company_name:
            company_para = doc.add_paragraph()
            company_para.add_run(branding.company_name).bold = True
            company_para.add_run(f"\n{branding.company_address}")
            company_para.add_run(f"\nTél: {branding.company_phone}")
            company_para.add_run(f"\nEmail: {branding.company_email}")
        
        # Générer les sections
        for section in template['sections']:
            if section['required'] or section['type'] in data:
                doc.add_heading(section['title'], level=1)
                
                section_content = data.get(section['type'], "Contenu à compléter")
                if isinstance(section_content, dict):
                    for key, value in section_content.items():
                        para = doc.add_paragraph()
                        para.add_run(f"{key}: ").bold = True
                        para.add_run(str(value))
                elif isinstance(section_content, list):
                    for item in section_content:
                        doc.add_paragraph(str(item), style='List Bullet')
                else:
                    doc.add_paragraph(str(section_content))
        
        # Pied de page
        if branding.footer_text:
            footer_para = doc.add_paragraph(branding.footer_text)
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Sauvegarder le document
        doc.save(str(output_file))
        
        return str(output_file)
    
    def _generate_excel_report(self, template: Dict[str, Any], data: Dict[str, Any], branding: BrandingConfig) -> str:
        """Génère un rapport Excel"""
        report_id = f"report_{uuid.uuid4().hex[:8]}"
        output_file = self.output_dir / f"{report_id}.xlsx"
        
        # Créer le classeur Excel
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Rapport"
        
        # Styles
        header_font = Font(name=branding.font_family, size=16, bold=True, 
                          color=branding.primary_color.replace('#', ''))
        section_font = Font(name=branding.font_family, size=12, bold=True,
                           color=branding.secondary_color.replace('#', ''))
        
        row = 1
        
        # Titre du rapport
        ws.cell(row=row, column=1, value=template['name'])
        ws.cell(row=row, column=1).font = header_font
        row += 2
        
        # Informations de l'entreprise
        if branding.company_name:
            ws.cell(row=row, column=1, value=branding.company_name)
            ws.cell(row=row, column=1).font = Font(bold=True)
            row += 1
            ws.cell(row=row, column=1, value=branding.company_address)
            row += 1
            ws.cell(row=row, column=1, value=f"Tél: {branding.company_phone}")
            row += 1
            ws.cell(row=row, column=1, value=f"Email: {branding.company_email}")
            row += 3
        
        # Générer les sections
        for section in template['sections']:
            if section['required'] or section['type'] in data:
                ws.cell(row=row, column=1, value=section['title'])
                ws.cell(row=row, column=1).font = section_font
                row += 1
                
                section_content = data.get(section['type'], "Contenu à compléter")
                if isinstance(section_content, dict):
                    for key, value in section_content.items():
                        ws.cell(row=row, column=1, value=key)
                        ws.cell(row=row, column=1).font = Font(bold=True)
                        ws.cell(row=row, column=2, value=str(value))
                        row += 1
                elif isinstance(section_content, list):
                    for item in section_content:
                        ws.cell(row=row, column=1, value=f"• {item}")
                        row += 1
                else:
                    ws.cell(row=row, column=1, value=str(section_content))
                    row += 1
                
                row += 1
        
        # Ajuster la largeur des colonnes
        ws.column_dimensions['A'].width = 30
        ws.column_dimensions['B'].width = 50
        
        # Sauvegarder le classeur
        wb.save(str(output_file))
        
        return str(output_file)
    
    def add_interactive_elements(self, report_path: str, hyperlinks: List[Dict[str, str]]) -> str:
        """Ajoute des éléments interactifs (liens hypertexte) à un rapport"""
        # Cette méthode peut être étendue pour ajouter des liens hypertexte
        # selon le format du rapport (PDF, Word, Excel)
        return report_path
    
    def get_available_sectors(self) -> Dict[str, str]:
        """Retourne la liste des secteurs d'activité disponibles"""
        return self.sector_templates
    
    def validate_template(self, template_data: Dict[str, Any]) -> List[str]:
        """Valide un template et retourne les erreurs éventuelles"""
        errors = []
        
        required_fields = ['name', 'description', 'sector', 'template_type', 'sections']
        for field in required_fields:
            if field not in template_data:
                errors.append(f"Champ requis manquant: {field}")
        
        if 'sections' in template_data:
            for i, section in enumerate(template_data['sections']):
                if 'type' not in section:
                    errors.append(f"Section {i}: type manquant")
                if 'title' not in section:
                    errors.append(f"Section {i}: titre manquant")
        
        return errors

